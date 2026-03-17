from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import os
from content_config import get_docx_content, DISCLAIMER, ENDING, FOOTER

def create_template(platform='jr', title='今日新闻鉴别报告', timestamp=None):
    doc = Document()

    now = datetime.datetime.now()
    weekday = now.weekday()
    weekdays_in_chinese = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
    date_str = f"{now.month}月{now.day}日  {weekdays_in_chinese[weekday]}"
    
    # 添加日期
    date_para = doc.add_paragraph('📅 日期：' + date_str)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加分隔线
    doc.add_paragraph('─' * 30)
    
    if timestamp:
        image_cover = f'output/image_{platform}_{timestamp}.png'
        image_poster = f'output/image2_{timestamp}.png'
        output_path = f'output/jrbd_{timestamp}.docx'
    else:
        image_cover = f'output/image_{platform}.png'
        image_poster = 'output/image2.png'
        output_path = 'jrbd.docx'
    
    # 添加封面图片
    if os.path.exists(image_cover):
        doc.add_picture(image_cover, width=Inches(6))
    
    # 添加介绍语
    intro_text = get_docx_content(title)
    intro_para = doc.add_paragraph(intro_text)
    intro_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 添加详细海报
    if os.path.exists(image_poster):
        doc.add_picture(image_poster, width=Inches(6))
    
    # 添加免责声明
    disclaimer_para = doc.add_paragraph(DISCLAIMER)
    disclaimer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in disclaimer_para.runs:
        run.font.size = Pt(10)
    
    # 添加结束语
    ending_para = doc.add_paragraph(ENDING)
    ending_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加页脚
    footer_para = doc.add_paragraph(FOOTER)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(9)
    
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")
    return output_path

if __name__ == "__main__":
    create_template('jr', '测试新闻标题')
    print("文档创建成功！")
